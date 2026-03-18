"""
Generates IEEE Methodology Section as a formatted Word document.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(0.875)
    section.right_margin = Inches(0.875)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

# ── Style Configuration ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.0

# Heading 1 style (Section title)
h1 = doc.styles['Heading 1']
h1.font.name = 'Times New Roman'
h1.font.size = Pt(10)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0, 0, 0)
h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
h1.paragraph_format.space_before = Pt(12)
h1.paragraph_format.space_after = Pt(6)

# Heading 2 style (Subsection)
h2 = doc.styles['Heading 2']
h2.font.name = 'Times New Roman'
h2.font.size = Pt(10)
h2.font.bold = False
h2.font.italic = True
h2.font.color.rgb = RGBColor(0, 0, 0)
h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h2.paragraph_format.space_before = Pt(8)
h2.paragraph_format.space_after = Pt(4)


def add_section_title(text):
    doc.add_heading(text, level=1)


def add_subsection(text):
    doc.add_heading(text, level=2)


def add_body(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)


def add_equation(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True


def add_pseudocode_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.bold = True


def add_pseudocode_line(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)


def add_diagram_placeholder(title, description):
    """Add a labeled diagram placeholder box."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"[DIAGRAM: {title}]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)
    run.bold = True
    run.font.color.rgb = RGBColor(80, 80, 80)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(6)
    run2 = p2.add_run(description)
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(8)
    run2.italic = True
    run2.font.color.rgb = RGBColor(100, 100, 100)


# ═══════════════════════════════════════════════════
# SECTION III. METHODOLOGY
# ═══════════════════════════════════════════════════

add_section_title("III. METHODOLOGY")

# ── A ──
add_subsection("A. System Overview")
add_body(
    "The proposed system implements an AI-based real-time fitness monitoring framework "
    "leveraging computer vision and human pose estimation to analyze exercise movements "
    "from live webcam or uploaded video input. The pipeline comprises four stages: "
    "(1) video acquisition and preprocessing, (2) pose estimation and landmark extraction, "
    "(3) kinematic analysis including joint angle computation, exercise recognition, "
    "repetition counting, and posture evaluation, and (4) feedback delivery and performance "
    "analytics. Each frame undergoes sequential processing, enabling continuous monitoring "
    "with per-frame latency suitable for real-time interaction. The system supports push-ups, "
    "squats, sit-ups, planks, jumping jacks, and bicep curls."
)

# ── B ──
add_subsection("B. System Architecture")
add_body(
    "The system adopts a modular layered architecture. The Input Layer handles video "
    "acquisition via OpenCV. The Processing Layer performs pose detection using MediaPipe "
    "PoseLandmarker, extracting 33 anatomical landmarks per frame. The Analysis Layer "
    "encapsulates joint angle computation, exercise-specific recognition, frame-stabilized "
    "repetition counting, and threshold-based posture evaluation. The Output Layer delivers "
    "real-time visual overlay feedback, with a proposed voice feedback module and a proposed "
    "analytics dashboard designed for future integration. The Storage Layer persists session "
    "data through a backend REST API."
)
add_diagram_placeholder(
    "Fig. 1. Overall System Architecture",
    "Input Layer (Webcam/Video) \u2192 Processing Layer (OpenCV \u2192 MediaPipe PoseLandmarker \u2192 "
    "33-Point Landmark Extraction) \u2192 Analysis Layer (Joint Angle Computation \u2192 Exercise "
    "Recognition \u2192 Repetition Counting \u2192 Posture Evaluation) \u2192 Output Layer (Visual Feedback, "
    "Proposed Voice Feedback, Proposed Analytics Dashboard) \u2192 Storage Layer (REST API \u2192 Database)"
)

# ── C ──
add_subsection("C. Video Processing Pipeline")
add_body(
    "Frame acquisition uses the OpenCV VideoCapture interface, supporting live camera "
    "streams and pre-recorded files. Each frame is converted from BGR to RGB and "
    "encapsulated as a MediaPipe Image object. For uploaded videos, per-frame time "
    "interval is computed as tf = 1/FPS to maintain temporal accuracy in duration-based "
    "exercises. Frames are annotated with skeleton overlays and feedback text before "
    "JPEG encoding for streaming delivery via HTTP multipart responses."
)
add_diagram_placeholder(
    "Fig. 2. Data Flow Pipeline",
    "Video Source \u2192 OpenCV Capture \u2192 BGR-to-RGB Conversion \u2192 MediaPipe Image Wrap \u2192 "
    "PoseLandmarker Detection \u2192 Landmark Extraction \u2192 Skeleton Overlay \u2192 JPEG Encode and Stream"
)

# ── D ──
add_subsection("D. Human Pose Estimation")
add_body(
    "The system employs MediaPipe PoseLandmarker, a lightweight neural network optimized "
    "for real-time inference. The model is configured with detection, presence, and tracking "
    "confidence thresholds of 0.7. For each frame, it produces 33 three-dimensional landmarks "
    "spanning the full body. Each landmark li is represented as (xi, yi, zi, vi), where "
    "(xi, yi) are normalized image coordinates, zi encodes relative depth, and vi is "
    "visibility confidence in the range [0, 1]. Landmarks with vi < 0.3 are excluded from "
    "analysis."
)
add_diagram_placeholder(
    "Fig. 3. Pose Estimation Pipeline",
    "Input Frame (RGB Image) \u2192 MediaPipe PoseLandmarker \u2192 33 Anatomical Landmarks \u2192 "
    "Visibility Check (v \u2265 0.3): Pass \u2192 Valid Coordinates \u2192 Downstream Analysis; "
    "Fail \u2192 Landmark Excluded"
)

# ── E ──
add_subsection("E. Landmark Coordinate Processing")
add_body(
    "Landmarks are stored as normalized coordinates, providing inherent scale and position "
    "invariance. The system performs bilateral selection by computing aggregate visibility "
    "scores for left and right anatomical groups, selecting the side with higher cumulative "
    "visibility as the primary reference. This ensures robustness to partial occlusion and "
    "varying camera angles."
)

# ── F ──
add_subsection("F. Joint Angle Computation")
add_body(
    "Given three landmark points Pa, Pb (vertex), and Pc, directional vectors "
    "A = Pa - Pb and B = Pc - Pb are constructed. The joint angle \u03b8 at the vertex Pb "
    "is computed using the inverse cosine of the normalized dot product:"
)
add_equation("\u03b8 = cos\u207b\u00b9( (A \u00b7 B) / (|A| \u00d7 |B|) )")
add_body(
    "The cosine value is clipped to [-1, 1] for numerical stability. The system supports "
    "both 2D (atan2-based) and 3D angle computation modes."
)
add_diagram_placeholder(
    "Fig. 4. Joint Angle Computation Flow",
    "Point A (Shoulder) and Point C (Wrist) \u2192 Construct Vectors A = A\u2212B, B = C\u2212B "
    "relative to Point B (Elbow/Vertex) \u2192 Dot Product and Magnitudes \u2192 \u03b8 = cos\u207b\u00b9(A\u00b7B / |A||B|)"
)

# ── G ──
add_subsection("G. Exercise Recognition Logic")
add_body(
    "Recognition uses rule-based kinematic analysis per exercise. Body orientation is "
    "validated: vertical (|x_shoulder - x_hip| < 0.2) for standing exercises such as squats, "
    "jumping jacks, and bicep curls; horizontal (|y_shoulder - y_hip| < 0.25) for floor "
    "exercises such as push-ups, planks, and sit-ups. Each exercise module implements a "
    "dedicated state-update function processing landmarks against exercise-specific thresholds."
)

# ── H ──
add_subsection("H. Repetition Counting Algorithm")
add_body(
    "Repetition counting employs a frame-stabilized finite state machine (FSM) with "
    "hysteresis thresholds. The algorithm maintains UP and DOWN states, requiring N "
    "consecutive frames at each threshold to confirm transitions, preventing false "
    "counts from noise."
)

add_pseudocode_title("Algorithm 1: Repetition Counting")
pseudocode_1 = [
    "FUNCTION CountRepetitions(angle \u03b8, state S):",
    "    INPUT: joint angle \u03b8, exercise state S",
    "    CONSTANTS: \u03b8_down, \u03b8_up, N_frames",
    "",
    "    IF \u03b8 < \u03b8_down THEN",
    "        S.count_down \u2190 S.count_down + 1",
    "        S.count_up \u2190 0",
    "        IF S.count_down \u2265 N AND NOT S.is_down THEN",
    "            S.is_down \u2190 TRUE",
    "    ELSE IF \u03b8 > \u03b8_up THEN",
    "        S.count_up \u2190 S.count_up + 1",
    "        S.count_down \u2190 0",
    "        IF S.count_up \u2265 N AND S.is_down THEN",
    "            S.reps \u2190 S.reps + 1",
    "            S.is_down \u2190 FALSE",
    "    RETURN S",
]
for line in pseudocode_1:
    add_pseudocode_line(line)

add_body(
    "Exercise-specific thresholds: push-ups (\u03b8_down = 110\u00b0, \u03b8_up = 145\u00b0, N = 2), "
    "squats (\u03b8_down = 100\u00b0, \u03b8_up = 160\u00b0, N = 3), sit-ups (\u03b8_down = 90\u00b0, \u03b8_up = 150\u00b0, N = 3)."
)
add_diagram_placeholder(
    "Fig. 5. Repetition Counting State Machine",
    "UP_STATE \u2192 (if \u03b8 < \u03b8_down) \u2192 CONFIRMING_DOWN \u2192 (if frames \u2265 N) \u2192 DOWN_STATE "
    "\u2192 (if \u03b8 > \u03b8_up) \u2192 CONFIRMING_UP \u2192 (if frames \u2265 N, count++) \u2192 UP_STATE"
)

# ── I ──
add_subsection("I. Posture Evaluation Model")
add_body(
    "Posture correctness uses a multi-criteria threshold model. The posture correctness "
    "score Sp for a single repetition is defined as follows:"
)
add_equation(
    "Sp = 3,  if \u03b8_min \u2264 \u03b8_perfect  AND  |I| = 0\n"
    "Sp = 2,  if \u03b8_min \u2264 \u03b8_good     AND  |I| = 0\n"
    "Sp = 1,  if \u03b8_min \u2264 \u03b8_down\n"
    "Sp = 0,  otherwise"
)
add_body(
    "where \u03b8_min is the minimum angle during the repetition, \u03b8_perfect and \u03b8_good are "
    "quality thresholds (e.g., 90\u00b0 and 100\u00b0 for push-ups), and |I| is the count of "
    "detected form issues such as hip sagging or elbow flaring. A repetition is classified "
    "as correct when Sp \u2265 2 and |I| = 0. The overall session accuracy is computed as "
    "A = N_correct / N_total \u00d7 100%."
)

add_pseudocode_title("Algorithm 2: Posture Evaluation")
pseudocode_2 = [
    "FUNCTION EvaluatePosture(\u03b8_min, issues):",
    "    IF \u03b8_min \u2264 \u03b8_perfect AND |issues| = 0 THEN",
    "        RETURN score \u2190 3, label \u2190 \"Excellent\"",
    "    ELSE IF \u03b8_min \u2264 \u03b8_good AND |issues| = 0 THEN",
    "        RETURN score \u2190 2, label \u2190 \"Good\"",
    "    ELSE IF \u03b8_min \u2264 \u03b8_down THEN",
    "        RETURN score \u2190 1, label \u2190 \"Shallow\"",
    "    ELSE",
    "        RETURN score \u2190 0, label \u2190 \"Incomplete\"",
]
for line in pseudocode_2:
    add_pseudocode_line(line)

add_diagram_placeholder(
    "Fig. 6. Posture Evaluation Logic",
    "Frame Landmarks \u2192 Compute Body Alignment Angle + Track Min Joint Angle \u2192 "
    "Check Hip Sag (angle < 150\u00b0) / Hip Pike (angle > 200\u00b0) \u2192 Add Issues \u2192 "
    "Depth Assessment: \u03b8 \u2264 90\u00b0 = Score 3, \u03b8 \u2264 100\u00b0 = Score 2, \u03b8 \u2264 110\u00b0 = Score 1, else = Score 0"
)

# ── J ──
add_subsection("J. Proposed Real-Time Voice Feedback Module")
add_body(
    "A voice feedback module is designed as a proposed component operating through four "
    "stages: (1) event detection triggering on posture deviations or repetition completions, "
    "(2) message generation via priority queue with deduplication, (3) speech synthesis through "
    "a TTS engine, and (4) audio delivery to the client. The planned design includes latency "
    "gating (minimum interval of 2 seconds between consecutive outputs) and priority-based "
    "scheduling favoring safety-critical corrections."
)
add_diagram_placeholder(
    "Fig. 7. Proposed Voice Feedback Pipeline",
    "Posture Event Detected \u2192 Priority Queue and Deduplication \u2192 Message Template Selection "
    "\u2192 TTS Engine (Proposed) \u2192 Audio Output to Client; with Latency Gate (\u0394t \u2265 2s)"
)

# ── K ──
add_subsection("K. Proposed Performance Analytics Dashboard")
add_body(
    "A web-based dashboard is designed as a proposed module for longitudinal tracking. "
    "The pipeline encompasses metric computation (repetitions, accuracy, depth scores), "
    "data persistence via REST API, statistical analysis (moving averages, improvement "
    "ratios), and visualization through a planned React.js frontend rendering interactive "
    "charts and progress summaries."
)
add_diagram_placeholder(
    "Fig. 8. Proposed Analytics Data Pipeline",
    "Session Metrics \u2192 REST API \u2192 Database Storage \u2192 Statistical Analysis \u2192 "
    "Trend Computation \u2192 React.js Dashboard (Proposed) \u2192 Charts and Reports"
)

# ── L ──
add_subsection("L. Backend and Database Integration")
add_body(
    "The Django backend exposes endpoints for video upload, live streaming, and status "
    "polling. Uploads are analyzed frame-by-frame, returning aggregated JSON results. "
    "Live streaming uses StreamingHttpResponse with MJPEG encoding. Session data including "
    "repetition counts, quality scores, form issues, and timestamps is serialized and persisted."
)
add_diagram_placeholder(
    "Fig. 9. Backend and Database Interaction",
    "Client (Browser) \u2192 Upload/Stream \u2192 Django Views \u2192 Analysis Engine \u2192 "
    "JSON Serializer \u2192 Store to Database / Response to Client; MJPEG Stream to Client"
)

# ── M ──
add_subsection("M. System Workflow Integration")
add_body(
    "Upon session initialization, the video source opens and the exercise-specific FSM "
    "is configured. Per frame: capture, pose detection, angle computation, state update, "
    "posture assessment, visual overlay, and stream delivery. Upon termination, aggregated "
    "metrics are stored. Proposed voice and analytics modules are designed for seamless "
    "integration at the feedback stage."
)
add_diagram_placeholder(
    "Fig. 10. End-to-End System Workflow",
    "Initialize Session \u2192 Capture Frame \u2192 Detect Landmarks \u2192 Compute Angles \u2192 "
    "Update FSM \u2192 Evaluate Posture \u2192 Render Overlay \u2192 Stream to Client \u2192 "
    "Loop if Active, else Store Metrics. Proposed modules (Voice Feedback, Dashboard) "
    "integrate at Posture Evaluation and Metrics Storage stages."
)

# ── N ──
add_subsection("N. Computational Considerations")
add_body(
    "Real-time performance is achieved through several strategies. MediaPipe PoseLandmarker "
    "operates with sub-30ms latency on standard CPU hardware, enabling 30 FPS processing. "
    "Rule-based analysis runs in O(1) per frame. Frame stabilization introduces controlled "
    "latency of N_frames / FPS seconds per transition while eliminating false detections "
    "from landmark jitter. Memory overhead is minimized by maintaining only current-frame "
    "landmarks and compact state dictionaries."
)

# ── Save ──
output_path = r"c:\Users\8thma\Exercise-Detection\IEEE_Methodology_Section.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
